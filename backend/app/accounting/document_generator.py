import os
import logging
from datetime import datetime
from typing import Dict, Any, Optional, List
from uuid import UUID
import io

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Pt, Inches

from app.accounting.models import Company, Obligation, Payment
from app.accounting.services import get_company, get_obligations, get_payments

logger = logging.getLogger(__name__)

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "generated_forms")

os.makedirs(TEMPLATES_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

class DocumentGenerator:
    """Document generator for accounting forms."""
    
    def __init__(self):
        """Initialize the document generator."""
        self.templates = self._load_templates()
        
    def _load_templates(self) -> Dict[str, str]:
        """Load available templates."""
        templates = {}
        for filename in os.listdir(TEMPLATES_DIR):
            name, ext = os.path.splitext(filename)
            if ext.lower() in ['.pdf', '.docx']:
                templates[name] = os.path.join(TEMPLATES_DIR, filename)
        return templates
    
    def get_available_templates(self) -> List[str]:
        """Get a list of available templates."""
        return list(self.templates.keys())
    
    def generate_form(
        self, 
        template_name: str, 
        company_id: int, 
        period: Optional[str] = None
    ) -> Optional[str]:
        """
        Generate a form from a template.
        
        Args:
            template_name: Name of the template
            company_id: ID of the company
            period: Optional period (YYYY-MM)
            
        Returns:
            Path to the generated document
        """
        if template_name not in self.templates:
            logger.error(f"Template {template_name} not found")
            return None
            
        template_path = self.templates[template_name]
        _, ext = os.path.splitext(template_path)
        
        company = get_company(company_id)
        if not company:
            logger.error(f"Company {company_id} not found")
            return None
            
        period_date = None
        if period:
            try:
                period_date = datetime.strptime(period, "%Y-%m")
            except ValueError:
                logger.error(f"Invalid period format: {period}, expected YYYY-MM")
                return None
                
        data = self._generate_form_data(company, period_date)
        
        if ext.lower() == '.pdf':
            return self._generate_pdf(template_path, data, template_name, company_id, period)
        elif ext.lower() == '.docx':
            return self._generate_docx(template_path, data, template_name, company_id, period)
        else:
            logger.error(f"Unsupported file extension: {ext}")
            return None
            
    def _generate_form_data(
        self, 
        company: Company, 
        period_date: Optional[datetime] = None
    ) -> Dict[str, Any]:
        """
        Generate data for form fields.
        
        Args:
            company: Company model
            period_date: Optional period date
            
        Returns:
            Dictionary of form field data
        """
        data = {
            "company_name": company.name,
            "company_location": company.location,
            "is_zona_libre": "Sí" if company.is_zona_libre else "No",
            "date": datetime.now().strftime("%Y-%m-%d"),
        }
        
        if period_date: 
            period_str = period_date.strftime("%Y-%m")
            data["period"] = period_str
            data["period_year"] = period_date.strftime("%Y")
            data["period_month"] = period_date.strftime("%m")
            
            month_start = period_date.replace(day=1)
            if period_date.month == 12:
                next_month = datetime(period_date.year + 1, 1, 1)
            else:
                next_month = datetime(period_date.year, period_date.month + 1, 1)
            
            obligations = get_obligations(filters={"company_id": company.id})
            period_obligations = []
            
            for obligation in obligations:
                due_date = obligation.next_due_date
                if isinstance(due_date, str):
                    due_date = datetime.fromisoformat(due_date.replace("Z", "+00:00"))
                
                if month_start <= due_date < next_month:
                    period_obligations.append(obligation)
            
            total_amount = sum(o.amount or 0 for o in period_obligations)
            data["total_obligations"] = len(period_obligations)
            data["total_amount"] = total_amount
            
            payments = []
            for obligation in period_obligations:
                obligation_payments = get_payments(filters={"obligation_id": obligation.id})
                payments.extend(obligation_payments)
            
            total_paid = sum(p.amount for p in payments)
            data["total_paid"] = total_paid
            data["total_pending"] = total_amount - total_paid
        
        return data
        
    def _generate_pdf(
        self, 
        template_path: str, 
        data: Dict[str, Any],
        template_name: str,
        company_id: int,
        period: Optional[str] = None
    ) -> str:
        """
        Generate a PDF document from a template.
        
        Args:
            template_path: Path to the template
            data: Form field data
            template_name: Name of the template
            company_id: ID of the company
            period: Optional period (YYYY-MM)
            
        Returns:
            Path to the generated PDF
        """
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        period_str = f"_{period}" if period else ""
        output_filename = f"{template_name}_company{company_id}{period_str}_{timestamp}.pdf"
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        try:
            packet = io.BytesIO()
            c = canvas.Canvas(packet, pagesize=letter)
            
            c.drawString(100, 700, f"Company: {data['company_name']}")
            c.drawString(100, 680, f"Location: {data['company_location']}")
            c.drawString(100, 660, f"Zona Libre: {data['is_zona_libre']}")
            c.drawString(100, 640, f"Date: {data['date']}")
            
            if 'period' in data:
                c.drawString(100, 620, f"Period: {data['period']}")
                c.drawString(100, 600, f"Total Obligations: {data['total_obligations']}")
                c.drawString(100, 580, f"Total Amount: ${data['total_amount']:.2f}")
                c.drawString(100, 560, f"Total Paid: ${data['total_paid']:.2f}")
                c.drawString(100, 540, f"Total Pending: ${data['total_pending']:.2f}")
            
            c.save()
            
            packet.seek(0)
            
            with open(output_path, 'wb') as f:
                f.write(packet.getvalue())
                
            logger.info(f"Generated PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            return None
        
    def _generate_docx(
        self, 
        template_path: str, 
        data: Dict[str, Any],
        template_name: str,
        company_id: int,
        period: Optional[str] = None
    ) -> str:
        """
        Generate a DOCX document from a template.
        
        Args:
            template_path: Path to the template
            data: Form field data
            template_name: Name of the template
            company_id: ID of the company
            period: Optional period (YYYY-MM)
            
        Returns:
            Path to the generated DOCX
        """
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        period_str = f"_{period}" if period else ""
        output_filename = f"{template_name}_company{company_id}{period_str}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        try:
            doc = Document(template_path)
            
            for para in doc.paragraphs:
                for key, value in data.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, str(value))
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            placeholder = f"{{{key}}}"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))
            
            doc.save(output_path)
            
            logger.info(f"Generated DOCX: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            return None

    def get_template_path(self, template_name: str) -> Optional[str]:
        """
        Get the path to a template.
        
        Args:
            template_name: Name of the template
            
        Returns:
            Path to the template or None if not found
        """
        if template_name in self.templates:
            return self.templates[template_name]
        return None

document_generator = DocumentGenerator()
